import re
from docx import Document
from docx2pdf import convert

filename = "data/TROD_angine_docx.docx"

dictionary_word = {"identifiant_pharmacie": "Pharmacie du Bois plage en ré",
                   "nom_pharmacien": "Jean Chatenet",
                   "date_heure_test": "22/09/2023 à 14h30",
                   "nom_patient": "tnndd92",
                   "age_patient": "19",
                   "p_ordo": "Non",
                   "score_mac": "3",
                   "1score": "1",
                   "2score": "1",
                   "3score": "1",
                   "4score": "0",
                   "5score": "0",
                   "6score": "0",
                   "7score": "3",
                   "nom_test": "Pfizer TROD Angine",
                   "lot_trod": "AK56FD",
                   "date_trod": "04/2026",
                   "ecouvillon_utilise": "celui prévu dans la boite",
                   "resultattest": "POSITIF",
                   "orientation_med": "Oui",
                   "deliv_antibio": "Non",
                   "trait_sympto": "Non",
                   "control_int": "Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed non risus.",
                   "control_ext": "Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed non risus.",
                   "date_val": "01/01/2023"
                   }


def docx_replace_regex(doc_obj, regex, replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace)


def replace_in_word(filename, dictionary):
    doc = Document(filename)
    for word, replacement in dictionary.items():
        word_re = re.compile(word)
        docx_replace_regex(doc, word_re, replacement)
    doc.save('data/outputs/result2.docx')
    convert("data/outputs/result2.docx", "data/outputs/result2.pdf")
    return doc

# replace_in_word(filename, dictionary)
